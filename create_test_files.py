"""
One-time script to generate sample PDF, DOCX, and Excel files for testing
multi-format document ingestion. Not part of the actual chatbot pipeline.
"""

from pathlib import Path
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF

DOCS_DIR = Path("data/documents")
DOCS_DIR.mkdir(parents=True, exist_ok=True)

# --- Create a test DOCX file ---
doc = Document()
doc.add_heading("Company Benefits Overview", level=1)
doc.add_paragraph(
    "The company offers a comprehensive health insurance plan covering employees "
    "and their immediate family members. Coverage begins on the first day of employment."
)
doc.add_paragraph(
    "Employees are also eligible for a 401(k) retirement plan with company matching "
    "up to 5% of their annual salary, starting after 90 days of employment."
)
doc.save(DOCS_DIR / "benefits_overview.docx")
print("Created benefits_overview.docx")

# --- Create a test Excel file ---
wb = Workbook()
ws = wb.active
ws.title = "Employees"
ws.append(["Name", "Department", "Years of Service"])
ws.append(["Ali Raza", "Engineering", 3])
ws.append(["Sara Khan", "Marketing", 5])
ws.append(["Bilal Ahmed", "Finance", 2])
wb.save(DOCS_DIR / "employee_directory.xlsx")
print("Created employee_directory.xlsx")

# --- Create a test PDF file ---
pdf = FPDF()
pdf.add_page()
pdf.set_font("Helvetica", size=12)
pdf.multi_cell(0, 10, (
    "IT Security Policy\n\n"
    "All employees must use two-factor authentication for company accounts. "
    "Passwords must be changed every 90 days and cannot be reused from the "
    "previous 5 passwords. Any suspected security breach must be reported to "
    "the IT department within 24 hours."
))
pdf.output(DOCS_DIR / "security_policy.pdf")
print("Created security_policy.pdf")

print("\nAll test files created in data/documents/")